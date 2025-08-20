import os
import json
import fitz  # PyMuPDF
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_cors import CORS
from docx import Document
import tempfile
import traceback

# --- App Configuration ---
app = Flask(__name__, static_folder='../client/dist', static_url_path='/')
CORS(app)

# --- Helper Functions ---

def extract_words_from_pdf(pdf_path, page_no=0):
    # ... (This function is good, no changes needed)
    doc = fitz.open(pdf_path)
    page = doc[page_no]
    words = page.get_text("words")
    spans = []
    for w in words:
        spans.append({
            "x0": w[0], "y0": w[1], "x1": w[2], "y1": w[3],
            "text": w[4], "cx": (w[0] + w[2]) / 2, "cy": (w[1] + w[3]) / 2
        })
    return spans

def find_nearby_text(spans, x, y, maxdist=50):
    # ... (This function is good, no changes needed)
    nearby = []
    for span in spans:
        dx = span["cx"] - x; dy = span["cy"] - y
        dist = (dx*dx + dy*dy)**0.5
        if dist <= maxdist:
            nearby.append(span["text"])
    return nearby

def docx_replace(doc, replacements):
    """
    A robust function to find and replace text placeholders in a .docx file,
    preserving formatting.
    """
    for key, value in replacements.items():
        placeholder = f"{{{{{key}}}}}"
        # Replace in all paragraphs
        for p in doc.paragraphs:
            if placeholder in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if placeholder in inline[i].text:
                        inline[i].text = inline[i].text.replace(placeholder, str(value))
        # Replace in all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if placeholder in p.text:
                            inline = p.runs
                            for i in range(len(inline)):
                                if placeholder in inline[i].text:
                                    inline[i].text = inline[i].text.replace(placeholder, str(value))
    return doc


# --- API Endpoints ---

@app.route('/api/resolve-balloon', methods=['POST'])
def resolve_balloon():
    # ... (This endpoint is good, no changes needed)
    try:
        if 'pdf' not in request.files: return jsonify({'error': 'No PDF file provided'}), 400
        pdf_file = request.files['pdf']; x = float(request.form['x']); y = float(request.form['y'])
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            pdf_path = tmp.name; pdf_file.save(pdf_path)
        spans = extract_words_from_pdf(pdf_path)
        nearby = find_nearby_text(spans, x, y)
        os.unlink(pdf_path)
        return jsonify({'nearby': nearby})
    except Exception as e:
        print(f"[ERROR] in /api/resolve-balloon: {traceback.format_exc()}"); return jsonify({'error': str(e)}), 500

@app.route('/api/generate-report', methods=['POST'])
def generate_report():
    output_filename = None
    try:
        print("\n--- [INFO] Received /api/generate-report request ---")
        
        # --- 1. Get Data ---
        balloons_json = request.form.get('balloons')
        if not balloons_json:
            print("[ERROR] No balloon data provided in the form.")
            return jsonify({"error": "No balloon data provided"}), 400
        
        balloon_data = json.loads(balloons_json)
        print(f"[DEBUG] Received {len(balloon_data)} balloons: {json.dumps(balloon_data, indent=2)}")

        # Hardcoded metadata for simplicity. Can be received from frontend later.
        metadata = {
            'customer_name': 'POLARIS', 'part_number': '5419070', 'report_date': '04/04/2025',
            'quantity': str(len(balloon_data)), 'prepared_by': 'Sudesh', 'verified_by': 'Dr. Sushant Maity'
        }
        print(f"[DEBUG] Metadata for report: {metadata}")

        # --- 2. Load Template ---
        template_path = "report_template.docx"
        if not os.path.exists(template_path):
             print(f"[ERROR] Template file not found at: {template_path}")
             return jsonify({"error": "Template file not found on server"}), 500
        
        doc = Document(template_path)
        print(f"[INFO] Loaded template: {template_path}")

        # --- 3. Replace Header/Footer Placeholders ---
        doc = docx_replace(doc, metadata)
        print("[INFO] Replaced header/footer metadata placeholders.")

        # --- 4. Populate the Dynamic Table ---
        if not doc.tables:
            print("[ERROR] No tables found in the document template.")
            return jsonify({"error": "Template is missing a table."}), 500
            
        table = doc.tables[0]
        template_row = table.rows[1] # The row with placeholders (e.g., {{id}})
        
        print(f"[INFO] Found table. Template row has {len(template_row.cells)} cells.")

        for balloon in balloon_data:
            new_row = table.add_row()
            # Copy cells from template to new row to preserve structure
            for i, cell in enumerate(template_row.cells):
                new_row.cells[i].text = cell.text
            
            # Create a dictionary of replacements for this specific balloon
            replacements = {
                'id': balloon.get('id', ''),
                'text': balloon.get('text', 'N/A'),
                'inspection_method': "Checking Fixture", # Example static value
                'remarks': balloon.get('type', '')
            }
            
            # Replace placeholders in the newly created row
            for cell in new_row.cells:
                for key, value in replacements.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                       # This simple replace is fine here since we just copied plain text
                       cell.text = cell.text.replace(placeholder, str(value))
        
        # Delete the original template row
        tbl = table._tbl
        tr_to_remove = template_row._tr
        tbl.remove(tr_to_remove)
        print(f"[INFO] Populated table with {len(balloon_data)} new rows and removed template row.")

        # --- 5. Save and Send ---
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            output_filename = tmp.name
            doc.save(output_filename)
            print(f"[INFO] Saved filled document to temporary file: {output_filename}")

        return send_file(
            output_filename,
            as_attachment=True,
            download_name='generated_report.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"[FATAL ERROR] in /api/generate-report: {traceback.format_exc()}")
        return jsonify({"error": "An internal server error occurred"}), 500
    finally:
        if output_filename and os.path.exists(output_filename):
            os.remove(output_filename)
            print(f"[INFO] Cleaned up temporary file: {output_filename}")
        print("--- [INFO] Request finished. ---\n")

# --- Serving the Frontend ---
@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def serve(path):
    # ... (This part is good, no changes needed)
    if path != "" and os.path.exists(os.path.join(app.static_folder, path)):
        return send_from_directory(app.static_folder, path)
    else:
        return send_from_directory(app.static_folder, 'index.html')

# --- Main Execution ---
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)