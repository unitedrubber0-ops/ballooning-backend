# server/docx_fill.py
from pathlib import Path

try:
    # prefer python-docx-replace (handles placeholders spanning runs)
    from python_docx_replace import docx_replace
    HAVE_DOCX_REPLACE = True
except Exception:
    HAVE_DOCX_REPLACE = False

from docx import Document


def fill_docx_template(template_path, output_path, replacements: dict):
    """
    Fill placeholders in a .docx template and write to output_path.
    - If python-docx-replace is available, docx_replace will be used (handles complex runs).
    - Otherwise, fallback to a simple run-level replace (may not catch placeholders split across runs).
    Placeholders expected in template like ${field_name} or {{field_name}} — pass keys accordingly.
    """
    template_path = Path(template_path)
    output_path = Path(output_path)

    if HAVE_DOCX_REPLACE:
        doc = Document(str(template_path))
        # docx_replace accepts keyword args; pass replacements as kwargs
        # ensure keys are strings without braces if the library expects direct names
        # first try with ${name} pattern: docx_replace replaces placeholders by kwargs keys
        # docx_replace(document, key1=value1, ...)
        try:
            docx_replace(doc, **replacements)
            doc.save(str(output_path))
            return output_path
        except Exception as e:
            # fallback to plain replace if docx_replace fails
            print("docx_replace failed:", e)

    # Fallback: simple run-based replacement
    doc = Document(str(template_path))
    placeholder_keys = list(replacements.keys())

    def replace_in_paragraph(paragraph):
        for key in placeholder_keys:
            # Accept both ${key} and {{key}} and plain key
            for token in (f"${{{key}}}", f"{{{{{key}}}}}", key):
                if token in paragraph.text:
                    # naive approach: update runs
                    # iterate through runs and replace token occurrences inside each run
                    for run in paragraph.runs:
                        if token in run.text:
                            run.text = run.text.replace(token, str(replacements[key]))

    # replace in paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p)
    # replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    doc.save(str(output_path))
    return output_path
