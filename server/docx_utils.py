from docx import Document
import os
import subprocess
import re

# --- Conversion helpers ---

def convert_doc_to_docx(path, soffice_path=None):
    if path.lower().endswith('.doc'):
        out = os.path.splitext(path)[0] + '.docx'
        cmd = [soffice_path or 'soffice', '--headless', '--convert-to', 'docx', path, '--outdir', os.path.dirname(path)]
        subprocess.run(cmd, check=True)
        return out
    return path

# --- Extraction helpers ---

def extract_table_rows(docx_path):
    doc = Document(docx_path)
    rows = []
    for table in doc.tables:
        for r in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in r.cells))
    return rows

# --- Replacement utilities ---

def replace_placeholders_docx(docx_path, out_path, mapping, prefer_preserve_format=True):
    """
    Replace placeholders everywhere (tables + paragraphs). If prefer_preserve_format is True, we try to
    modify runs minimally; fallback: rebuild the paragraph text which may lose inline formatting in that paragraph.
    """
    doc = Document(docx_path)

    # Small helper to replace in a run-aware way
    def replace_in_paragraph_runs(paragraph, mapping):
        # Join runs into a single string with index mapping
        full_text = "".join(run.text for run in paragraph.runs)
        replaced = full_text
        for k, v in mapping.items():
            replaced = replaced.replace(k, str(v))
        if replaced == full_text:
            return
        if prefer_preserve_format and len(paragraph.runs) == 1:
            paragraph.runs[0].text = replaced
            return
        # Fallback: clear and add single run (preserves paragraph style but not inline formatting)
        for _ in range(len(paragraph.runs)):
            paragraph.runs[0].clear()
            paragraph.runs[0].text = ''
            paragraph.runs[0].element.getparent().remove(paragraph.runs[0].element)
        run = paragraph.add_run(replaced)

    # Replace in paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph_runs(p, mapping)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph_runs(p, mapping)

    doc.save(out_path)
    return out_path